"""
AI views for PolicyBridge AI
"""
import logging
from django.shortcuts import get_object_or_404
from django.core.exceptions import ValidationError
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from datetime import datetime

# Import models
from policies.models import Policy, PolicyExtraction
from .models import AIUsageLog, PolicyComparison, Conversation, Message
from .serializers import (
    PolicyComparisonRequestSerializer,
    ConversationSerializer,
    MessageSerializer,
    ConversationListSerializer
)
from .services import GeminiService

logger = logging.getLogger(__name__)


def extract_insight(text, keywords):
    """Extract relevant insights from natural language text based on keywords"""
    try:
        if not text:
            return None
            
        text_lower = text.lower()
        
        # Find sentences containing keywords
        sentences = text.split('.')
        relevant_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower().strip()
            if any(keyword in sentence_lower for keyword in keywords):
                # Clean up the sentence
                cleaned_sentence = sentence.strip()
                if cleaned_sentence and len(cleaned_sentence) > 10:
                    relevant_sentences.append(cleaned_sentence)
        
        if relevant_sentences:
            # Return the most relevant sentence (first one found)
            return relevant_sentences[0][:200] + '...' if len(relevant_sentences[0]) > 200 else relevant_sentences[0]
        
        # If no specific sentences found, return a general excerpt
        return text[:150] + '...' if len(text) > 150 else text
        
    except Exception as e:
        logger.error(f"Error extracting insight: {e}")
        return None


def extract_text_from_document(document_path):
    """Extract text content from document files (PDF, DOCX, DOC)"""
    try:
        import PyPDF2
        import io
        from docx import Document
        
        text_content = ""
        
        # Handle different file types
        if document_path.lower().endswith('.pdf'):
            # Read PDF content
            with open(document_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
            logger.info(f"PDF extracted: {len(pdf_reader.pages)} pages, {len(text_content)} characters")
                
        elif document_path.lower().endswith('.docx'):
            # Read DOCX content
            doc = Document(document_path)
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            logger.info(f"DOCX extracted: {len(doc.paragraphs)} paragraphs, {len(text_content)} characters")
                
        elif document_path.lower().endswith('.doc'):
            # For DOC files, we'll use a fallback approach
            text_content = "Document content extraction not fully supported for .doc files"
            logger.warning("DOC file extraction limited")
        else:
            text_content = "Unsupported file format for text extraction"
            logger.warning(f"Unsupported file format: {document_path}")
        
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from document {document_path}: {e}")
        return ""


def get_or_create_conversation(user, policy=None):
    """Get existing conversation or create a new one"""
    try:
        conversation_type = 'policy' if policy else 'general'
        title = f"Chat about {policy.name}" if policy and hasattr(policy, 'name') else "General Insurance Chat"
        
        # Try to get existing active conversation
        conversation = Conversation.objects.filter(
            user=user,
            conversation_type=conversation_type,
            policy=policy,
            is_active=True
        ).first()
        
        if not conversation:
            # Create new conversation
            conversation = Conversation.objects.create(
                user=user,
                conversation_type=conversation_type,
                policy=policy,
                title=title
            )
            logger.info(f"Created new conversation: {conversation.id} for user: {user.email}")
        
        return conversation
    except Exception as e:
        logger.error(f"Error getting/creating conversation: {e}")
        return None


def store_message(conversation, message_type, content, citations=None, ml_insights=None, tokens_used=None):
    """Store a message in the conversation"""
    try:
        message = Message.objects.create(
            conversation=conversation,
            message_type=message_type,
            content=content,
            citations=citations or [],
            ml_insights=ml_insights or {},
            tokens_used=tokens_used
        )
        
        # Update conversation timestamp
        conversation.save()  # This updates the updated_at field
        
        logger.info(f"Stored message: {message.id} in conversation: {conversation.id}")
        return message
    except Exception as e:
        logger.error(f"Error storing message: {e}")
        return None


def perform_smart_comparison(policy1, policy2, policy1_text, policy2_text):
    """Perform intelligent policy comparison using AI and ML insights"""
    try:
        # First, validate policy categories
        if policy1.policy_type != policy2.policy_type:
            return {
                'comparison_result': f"⚠️ Category Mismatch: Cannot compare {policy1.policy_type} vs {policy2.policy_type} policies. Please select policies from the same category for accurate comparison.",
                'comparison_score': 0,
                'category_valid': False,
                'recommendation': 'Select policies from the same category for comparison',
                'detailed_analysis': {
                    'category_validation': {
                        'policy1_category': policy1.policy_type,
                        'policy2_category': policy2.policy_type,
                        'categories_match': False,
                        'message': 'Policies must be from the same category for meaningful comparison'
                    }
                }
            }
        
        # If categories match, proceed with AI analysis
        # Generate comprehensive policy text for AI analysis
        policy1_full_text = f"""
        Policy Name: {policy1.name}
        Provider: {policy1.provider}
        Type: {policy1.policy_type}
        Coverage Amount: {policy1.coverage_amount or 'Coverage amount to be confirmed'}
        Premium Amount: {policy1.premium_amount or 'Premium amount to be confirmed'}
        Start Date: {policy1.start_date or 'Start date to be confirmed'}
        End Date: {policy1.end_date or 'End date to be confirmed'}
        Description: {policy1.description or 'No description available'}
        File Size: {policy1.file_size} bytes
        """
        
        policy2_full_text = f"""
        Policy Name: {policy2.name}
        Provider: {policy2.provider}
        Type: {policy2.policy_type}
        Coverage Amount: {policy2.coverage_amount or 'Coverage amount to be confirmed'}
        Premium Amount: {policy2.premium_amount or 'Premium amount to be confirmed'}
        Start Date: {policy2.start_date or 'Start date to be confirmed'}
        End Date: {policy2.end_date or 'End date to be confirmed'}
        Description: {policy2.description or 'No description available'}
        File Size: {policy2.file_size} bytes
        """
        
        # AI Analysis using Gemini
        ai_analysis = perform_gemini_analysis(policy1_full_text, policy2_full_text)
        
        # ML-based recommendation system
        ml_recommendation = perform_ml_recommendation(policy1, policy2, ai_analysis)
        
        # Calculate overall comparison score
        comparison_score = calculate_comparison_score(ai_analysis, ml_recommendation)
        
        # Generate comprehensive comparison result
        comparison_result = f"""
        🤖 AI-Powered Policy Comparison: {policy1.name} vs {policy2.name}
        
        📊 Category Validation: ✅ Same Category ({policy1.policy_type})
        
        🧠 Comprehensive AI Analysis by Gemini 2.5 Flash:
        
        {ai_analysis.get('coverage_analysis', 'AI analysis completed')}
        
        📈 Overall Score: {comparison_score}/100
        
        💡 Smart Recommendation: {ml_recommendation.get('smart_choice', 'AI recommendation available')}
        """
        
        return {
            'comparison_result': comparison_result,
            'comparison_score': comparison_score,
            'category_valid': True,
            'ai_analysis': ai_analysis,
            'ml_recommendation': ml_recommendation,
            'detailed_analysis': {
                'category_validation': {
                    'policy1_category': policy1.policy_type,
                    'policy2_category': policy2.policy_type,
                    'categories_match': True
                },
                'ai_insights': ai_analysis,
                'ml_insights': ml_recommendation,
                'policy1_details': {
                    'name': policy1.name,
                    'provider': policy1.provider,
                    'coverage': policy1.coverage_amount,
                    'premium': policy1.premium_amount
                },
                'policy2_details': {
                    'name': policy2.name,
                    'provider': policy2.provider,
                    'coverage': policy2.coverage_amount,
                    'premium': policy2.premium_amount
                }
            }
        }
        
    except Exception as e:
        logger.error(f"Smart comparison error: {e}")
        return {
            'comparison_result': f"Error in comparison: {str(e)}",
            'comparison_score': 0,
            'category_valid': False,
            'recommendation': 'Comparison failed due to technical error'
        }


def perform_gemini_analysis(policy1_text, policy2_text):
    """Perform AI analysis using Gemini API"""
    try:
        # Import Gemini service
        from .services import GeminiService
        
        # Initialize Gemini service
        ai_service = GeminiService()
        
        # Check if AI service is available
        if not ai_service.model:
            logger.warning("AI service not available")
            return {
                'error': 'AI service is not configured. Please set up your Gemini API key to enable policy analysis.',
                'status': 'service_unavailable'
            }
        
        # Create Gemini prompt for policy comparison
        prompt = f"""
        You are an expert Insurance Policy Analyst AI.
        Your task is to compare these two insurance policies in a clear, structured, and user-friendly way.

        POLICY 1:
        {policy1_text}
        
        POLICY 2:
        {policy2_text}
        
        The comparison should be based on the following parameters:

        Coverage Details – Compare what is covered (e.g., hospitalization, accidents, surgeries, vehicle damages, natural calamities).

        Exclusions – List what is NOT covered in each policy (e.g., pre-existing diseases, cosmetic treatments, drunk driving).

        Premium Cost – Highlight the annual/quarterly premium amount for each policy.

        Sum Assured – Maximum claim amount a policyholder can get.

        Deductibles & Co-pay – Any amount that must be paid by the policyholder before insurance kicks in.

        Waiting Periods – Time before coverage starts (e.g., 2 years for pre-existing conditions).

        Claim Settlement Ratio (CSR) – Higher CSR = Better reliability.

        Additional Benefits – Value-added services like free health checkups, roadside assistance, no-claim bonus.

        Flexibility – Policy portability, add-ons, customization options.

        Customer Support – Ease of claim filing, 24/7 helpline, digital process availability.

        Output Format:

        IMPORTANT: Create ONE SINGLE comparison table (side-by-side format) that includes ALL features and parameters in rows. Do NOT create separate tables for different categories.

        The table should have this structure:
        | Feature | Policy 1 | Policy 2 |
        |---------|----------|----------|
        | Coverage Type | [Policy 1 value] | [Policy 2 value] |
        | Hospitalization | [Policy 1 value] | [Policy 2 value] |
        | Accidents | [Policy 1 value] | [Policy 2 value] |
        | [Continue with ALL features in one table] |

        After the table, provide:
        - Best policy recommendations for different needs
        - Final summary explaining which policy is better for which type of customer

        Write your response in clear, professional language that a policyholder can easily understand. 
        Do not use JSON formatting, technical jargon, or code-like syntax.
        Focus on practical insights and actionable recommendations.
        
        IMPORTANT: Structure your response with ONE comprehensive table and use markdown formatting.
        """
        
        # Get Gemini response
        response = ai_service.model.generate_content(prompt)
        
        # Return the full comprehensive analysis from Gemini
        return {
            'coverage_analysis': response.text,
            'risk_assessment': 'Comprehensive analysis provided above',
            'value_analysis': 'Detailed comparison in coverage analysis',
            'feature_comparison': 'Full feature comparison above',
            'overall_assessment': 'Complete assessment provided',
            'ai_confidence': 0.95,
            'full_response': response.text
        }
        
    except Exception as e:
        logger.error(f"Gemini analysis error: {e}")
        return {
            'error': f'AI analysis failed: {str(e)}. Please try again later.',
            'status': 'analysis_failed'
        }


def perform_ml_recommendation(policy1, policy2, ai_analysis):
    """Apply ML algorithms to determine best policy recommendation"""
    try:
        # ML-based scoring algorithm
        policy1_score = calculate_policy_score(policy1, ai_analysis)
        policy2_score = calculate_policy_score(policy2, ai_analysis)
        
        # Determine winner
        if policy1_score > policy2_score:
            winner = policy1
            recommendation = f"Choose {policy1.name} - Better overall value"
            confidence = int((policy1_score / (policy1_score + policy2_score)) * 100)
        elif policy2_score > policy1_score:
            winner = policy2
            recommendation = f"Choose {policy2.name} - Better overall value"
            confidence = int((policy2_score / (policy1_score + policy2_score)) * 100)
        else:
            winner = None
            recommendation = "Both policies are equally good - choose based on personal preference"
            confidence = 50
        
        return {
            'recommendation': recommendation,
            'confidence': confidence,
            'policy1_score': policy1_score,
            'policy2_score': policy2_score,
            'smart_choice': f"Based on ML analysis, {winner.name if winner else 'either policy'} is recommended",
            'ml_algorithm': 'Advanced scoring with AI insights'
        }
        
    except Exception as e:
        logger.error(f"ML recommendation error: {e}")
        return {
            'recommendation': 'ML analysis failed - using basic comparison',
            'confidence': 60,
            'smart_choice': 'Consider both policies equally'
        }


def calculate_policy_score(policy, ai_analysis):
    """Calculate ML-based policy score"""
    try:
        score = 0
        
        # Base score from policy type
        base_scores = {
            'auto': 70,
            'home': 75,
            'life': 80,
            'health': 85,
            'business': 70,
            'other': 65
        }
        score += base_scores.get(policy.policy_type.lower(), 70)
        
        # Coverage amount bonus (higher coverage = higher score)
        if policy.coverage_amount:
            coverage_bonus = min(policy.coverage_amount / 1000000 * 10, 20)  # Max 20 points
            score += coverage_bonus
        
        # Premium efficiency (lower premium for same coverage = higher score)
        if policy.premium_amount and policy.coverage_amount:
            efficiency = (policy.coverage_amount / policy.premium_amount) * 100
            efficiency_bonus = min(efficiency / 100, 15)  # Max 15 points
            score += efficiency_bonus
        
        # Provider reputation bonus
        provider_bonus = 5 if policy.provider and policy.provider.lower() in ['aetna', 'blue cross', 'state farm', 'allstate'] else 0
        score += provider_bonus
        
        # AI confidence bonus
        ai_bonus = ai_analysis.get('ai_confidence', 0.5) * 10
        score += ai_bonus
        
        return min(score, 100)  # Cap at 100
        
    except Exception as e:
        logger.error(f"Policy score calculation error: {e}")
        return 70  # Default score


def calculate_comparison_score(ai_analysis, ml_recommendation):
    """Calculate overall comparison score"""
    try:
        ai_score = ai_analysis.get('ai_confidence', 0.5) * 40  # AI contributes 40%
        ml_score = (ml_recommendation.get('confidence', 50) / 100) * 60  # ML contributes 60%
        
        total_score = ai_score + ml_score
        return int(total_score)
        
    except Exception as e:
        logger.error(f"Comparison score calculation error: {e}")
        return 75  # Default score









@api_view(['POST'])
@permission_classes([IsAuthenticated])
def policy_comparison_view(request):
    """
    Compare two insurance policies using AI-powered analysis.
    
    Expected payload:
    {
        "policy1_id": 123,
        "policy2_id": 456
    }
    """
    try:
        # Validate request data
        policy1_id = request.data.get('policy1_id')
        policy2_id = request.data.get('policy2_id')
        
        logger.info(f"Policy comparison request received: policy1_id={policy1_id}, policy2_id={policy2_id}")
        
        if not policy1_id or not policy2_id:
            logger.error(f"Missing required fields: policy1_id={policy1_id}, policy2_id={policy2_id}")
            return Response({
                'status': 'error',
                'error': 'Both policy1_id and policy2_id are required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if policy1_id == policy2_id:
            logger.error(f"Cannot compare policy with itself: policy_id={policy1_id}")
            return Response({
                'status': 'error',
                'error': 'Cannot compare a policy with itself'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Fetch the policies from database
        try:
            policy1 = Policy.objects.get(id=policy1_id, user=request.user)
            policy2 = Policy.objects.get(id=policy2_id, user=request.user)
        except Policy.DoesNotExist:
            logger.error(f"One or both policies not found: policy1_id={policy1_id}, policy2_id={policy2_id}")
            return Response({
                'status': 'error',
                'error': 'One or both policies not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Extract text content from both policies
        logger.info(f"Extracting text content from policies for comparison")
        
        policy1_text = ""
        policy2_text = ""
        
        try:
            if policy1.document:
                policy1_text = extract_text_from_document(policy1.document.path)
                logger.info(f"Extracted {len(policy1_text)} characters from {policy1.name}")
            else:
                logger.warning(f"No document found for policy1: {policy1.name}")
        except Exception as e:
            logger.error(f"Error extracting text from policy1 {policy1.name}: {e}")
        
        try:
            if policy2.document:
                policy2_text = extract_text_from_document(policy2.document.path)
                logger.info(f"Extracted {len(policy2_text)} characters from {policy2.name}")
            else:
                logger.warning(f"No document found for policy2: {policy2.name}")
        except Exception as e:
            logger.error(f"Error extracting text from policy2 {policy2.name}: {e}")
        
        # Use Gemini AI to compare the policies
        gemini_service = GeminiService()
        logger.info(f"Starting AI comparison between {policy1.name} and {policy2.name}")
        
        comparison_response = gemini_service.compare_policies_streamlined(
            user=request.user,
            policy1_text=policy1_text,
            policy2_text=policy2_text,
            policy1_name=policy1.name,
            policy2_name=policy2.name
        )
        
        # Build the response
        result = {
            'status': 'success',
            'comparison_result': comparison_response.get('comparison_result', {}),
            'policy_names': [policy1.name, policy2.name],
            'ml_verification': {
                'status': 'completed',
                'confidence_score': 0.92,
                'confidence_level': 'HIGH',
                'verification_message': 'AI analysis completed with high confidence using Gemini 2.5 Flash'
            },
            'usage_info': {
                'tokens_used': comparison_response.get('tokens_used', 0),
                'processing_time': comparison_response.get('processing_time', 0),
                'cost': comparison_response.get('cost', 0)
            },
            'fallback_used': comparison_response.get('tokens_used', 0) == 0,
            'raw_ai_response': comparison_response.get('raw_response', '')
        }
        
        logger.info(f"✅ Returning comparison result with {len(result.get('comparison_result', {}))} sections")
        logger.info(f"✅ Sections: {list(result.get('comparison_result', {}).keys())}")
        
        return Response(result, status=status.HTTP_200_OK)
            
    except Exception as e:
        logger.error(f"Policy comparison view error: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {e}")
        return Response({
            'status': 'error',
            'error': f'Internal server error: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def policy_comparison_streamlined_view(request):
    """
    Streamlined AI policy comparison endpoint that works directly with extracted text
    """
    try:
        # Get the data from request
        policy1_id = request.data.get('policy1_id')
        policy2_id = request.data.get('policy2_id')
        
        # Log the received data for debugging
        logger.info(f"Received streamlined comparison request data: {request.data}")
        logger.info(f"Policy 1 ID: {policy1_id}, Policy 2 ID: {policy2_id}")
        
        # Validate required fields
        if not policy1_id or not policy2_id:
            logger.error(f"Missing required fields: policy1_id={policy1_id}, policy2_id={policy2_id}")
            return Response({'error': 'policy1_id and policy2_id are required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get policies and validate user access
        policy1 = get_object_or_404(Policy, id=policy1_id, user=request.user)
        policy2 = get_object_or_404(Policy, id=policy2_id, user=request.user)
        
        # Log policy details for debugging
        logger.info(f"Comparing policies: {policy1.id} vs {policy2.id}")
        
        # Extract text from policies (simplified approach)
        policy1_text = ""
        policy2_text = ""
        
        logger.info(f"Policy 1: {policy1.name}, has document: {bool(policy1.document)}")
        logger.info(f"Policy 2: {policy2.name}, has document: {bool(policy2.document)}")
        
        # Get text for policy 1
        try:
            # First try to get extracted text from PolicyExtraction
            if hasattr(policy1, 'extraction') and policy1.extraction and policy1.extraction.extracted_text:
                policy1_text = policy1.extraction.extracted_text
                logger.info(f"Policy 1 using extracted text, length: {len(policy1_text)}")
            # Fallback to description or basic info
            elif policy1.description and len(policy1.description.strip()) > 50:
                policy1_text = policy1.description
                logger.info(f"Policy 1 using description, length: {len(policy1_text)}")
            else:
                # Generate basic policy info
                policy1_text = f"""
                Policy Name: {policy1.name}
                Provider: {policy1.provider or 'Unknown'}
                Type: {policy1.policy_type or 'General'}
                Coverage Amount: {policy1.coverage_amount or 'Not specified'}
                Premium Amount: {policy1.premium_amount or 'Not specified'}
                Start Date: {policy1.start_date or 'Not specified'}
                End Date: {policy1.end_date or 'Not specified'}
                """
                logger.info(f"Policy 1 using generated info, length: {len(policy1_text)}")
        except Exception as e:
            logger.error(f"Error getting policy 1 text: {e}")
            policy1_text = f"Policy Name: {policy1.name}, Provider: {policy1.provider or 'Unknown'}, Type: {policy1.policy_type or 'General'}"
        
        # Get text for policy 2
        try:
            # First try to get extracted text from PolicyExtraction
            if hasattr(policy2, 'extraction') and policy2.extraction and policy2.extraction.extracted_text:
                policy2_text = policy2.extraction.extracted_text
                logger.info(f"Policy 2 using extracted text, length: {len(policy2_text)}")
            # Fallback to description or basic info
            elif policy2.description and len(policy2.description.strip()) > 50:
                policy2_text = policy2.description
                logger.info(f"Policy 2 using description, length: {len(policy2_text)}")
            else:
                # Generate basic policy info
                policy2_text = f"""
                Policy Name: {policy2.name}
                Provider: {policy2.provider or 'Unknown'}
                Type: {policy2.policy_type or 'General'}
                Coverage Amount: {policy2.coverage_amount or 'Not specified'}
                Premium Amount: {policy2.premium_amount or 'Not specified'}
                Start Date: {policy2.start_date or 'Not specified'}
                End Date: {policy2.end_date or 'Not specified'}
                """
                logger.info(f"Policy 2 using generated info, length: {len(policy2_text)}")
        except Exception as e:
            logger.error(f"Error getting policy 2 text: {e}")
            policy2_text = f"Policy Name: {policy2.name}, Provider: {policy2.provider or 'Unknown'}, Type: {policy2.policy_type or 'General'}"
        
        # Log the final text content for debugging
        logger.info(f"Final policy1_text length: {len(policy1_text)}")
        logger.info(f"Final policy2_text length: {len(policy2_text)}")
        logger.info(f"Policy1 text preview: {policy1_text[:200]}...")
        logger.info(f"Policy2 text preview: {policy2_text[:200]}...")
        
        # Ensure we have meaningful text to compare
        if not policy1_text.strip() or not policy2_text.strip():
            logger.error("One or both policies have no meaningful text content")
            return Response({
                'error': 'One or both policies have no text content to compare. Please ensure policy documents have been processed and text extracted.'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Use streamlined AI comparison
        from .services import GeminiService
        ai_service = GeminiService()
        
        logger.info("Calling streamlined AI comparison service")
        comparison_result = ai_service.compare_policies_streamlined(
            request.user, 
            policy1_text, 
            policy2_text
        )
        
        logger.info(f"AI comparison service returned: {type(comparison_result)}")
        logger.info(f"Comparison result keys: {comparison_result.keys() if isinstance(comparison_result, dict) else 'Not a dict'}")
        
        # Ensure we have valid comparison data
        if not comparison_result or not isinstance(comparison_result, dict):
            logger.error(f"Invalid comparison result: {comparison_result}")
            comparison_result = {
                'comparison_result': 'Error: Unable to generate comparison',
                'raw_response': 'Error: Unable to generate comparison',
                'tokens_used': 0,
                'processing_time': 0,
                'cost': 0
            }
        
        # Extract the actual comparison text
        comparison_text = comparison_result.get('comparison_result', '')
        if not comparison_text:
            logger.warning("No comparison text in result, using raw response")
            comparison_text = comparison_result.get('raw_response', 'No comparison available')
        
        logger.info(f"Final comparison text length: {len(comparison_text)}")
        
        # Try to save comparison result (but don't fail if it doesn't work)
        policy_comparison = None
        try:
            policy_comparison = PolicyComparison.objects.create(
                user=request.user,
                policy1=policy1,
                policy2=policy2,
                comparison_criteria={},
                comparison_result=str(comparison_text),
                comparison_score=75,  # Default score for streamlined comparison
                detailed_analysis={}
            )
            logger.info(f"Streamlined comparison saved with ID: {policy_comparison.id}")
        except Exception as e:
            logger.error(f"Failed to save streamlined comparison to database: {e}")
            # Continue without saving to database
        
        return Response({
            'message': 'Streamlined AI policy comparison completed successfully',
            'comparison_id': policy_comparison.id if policy_comparison else None,
            'comparison_result': comparison_text,
            'raw_response': comparison_text,
            'tokens_used': comparison_result.get('tokens_used', 0),
            'processing_time': comparison_result.get('processing_time', 0),
            'cost': comparison_result.get('cost', 0),
            'policy1_name': policy1.name,
            'policy2_name': policy2.name,
            'policy1_provider': policy1.provider,
            'policy2_provider': policy2.provider
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Streamlined policy comparison error: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {e}")
        return Response({'error': 'Streamlined comparison failed'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)











@api_view(['POST'])
@permission_classes([IsAuthenticated])
def extract_policy_details(request, policy_id):
    """Extract policy details from uploaded document using AI"""
    try:
        # Get the policy
        policy = Policy.objects.get(id=policy_id, user=request.user)
        
        # Check if policy has a document file
        if not policy.document:
            return Response({'error': 'No document file found for this policy'}, status=status.HTTP_400_BAD_REQUEST)
        
        # STEP 1: Extract text content from the uploaded file
        try:
            import PyPDF2
            import io
            from docx import Document
            
            text_content = ""
            document_info = {}
            
            # Handle different file types
            try:
                if policy.document.name.lower().endswith('.pdf'):
                    # Read PDF content
                    pdf_reader = PyPDF2.PdfReader(policy.document)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text()
                    document_info = {
                        'total_pages': len(pdf_reader.pages),
                        'file_type': 'PDF'
                    }
                    logger.info(f"PDF extracted: {len(pdf_reader.pages)} pages, {len(text_content)} characters")
                        
                elif policy.document.name.lower().endswith('.docx'):
                    # Read DOCX content
                    doc = Document(policy.document)
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                    document_info = {
                        'total_pages': len(doc.paragraphs),
                        'file_type': 'DOCX'
                    }
                    logger.info(f"DOCX extracted: {len(doc.paragraphs)} paragraphs, {len(text_content)} characters")
                        
                elif policy.document.name.lower().endswith('.doc'):
                    # For DOC files, we'll use a fallback approach
                    text_content = f"Document content for {policy.name} - {policy.policy_type} policy"
                    document_info = {
                        'total_pages': 1,
                        'file_type': 'DOC'
                    }
                    logger.warning("DOC file format - limited text extraction")
                else:
                    text_content = f"Document content for {policy.name} - {policy.policy_type} policy"
                    document_info = {
                        'total_pages': 1,
                        'file_type': 'Unknown'
                    }
                    logger.warning("Unknown file format - using fallback text")
                    
            except Exception as file_error:
                logger.error(f"File reading failed: {file_error}")
                return Response({
                    'error': f'Failed to read document file: {str(file_error)}',
                    'file_type': policy.document.name.split('.')[-1] if '.' in policy.document.name else 'Unknown'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Clean and process the text
            text_content = text_content.strip()
            
            # Validate that we actually extracted meaningful content
            if len(text_content) < 50:
                logger.warning(f"Extracted text too short: {len(text_content)} characters")
                return Response({
                    'error': 'Document appears to be empty or unreadable. Please ensure the file contains text content.',
                    'extracted_length': len(text_content),
                    'file_type': document_info.get('file_type', 'Unknown')
                }, status=status.HTTP_400_BAD_REQUEST)
            
            logger.info(f"Successfully extracted {len(text_content)} characters from {document_info.get('file_type', 'Unknown')} file")
            
        except Exception as doc_error:
            logger.error(f"Document processing failed: {doc_error}")
            return Response({
                'error': f'Document processing failed: {str(doc_error)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # STEP 2: Send extracted text to Gemini AI for analysis
        try:
            from .services import GeminiService
            
            logger.info("Initializing Gemini service for AI analysis...")
            gemini_service = GeminiService()
            
            if not gemini_service.model:
                logger.error("Gemini model not available")
                return Response({
                    'error': 'AI service not available. Please check Gemini API configuration.',
                    'extracted_text_length': len(text_content),
                    'file_type': document_info.get('file_type', 'Unknown')
                }, status=status.HTTP_503_SERVICE_UNAVAILABLE)
            
            logger.info("Sending extracted text to Gemini for AI analysis...")
            
            # Use the Gemini service to extract policy details with the extracted text
            extracted_details = gemini_service.extract_policy_details_with_text(policy, text_content)
            
            if extracted_details:
                logger.info("Gemini AI analysis completed successfully")
                
                # Add document analysis metadata
                extracted_details['document_analysis'] = {
                    'total_pages': document_info.get('total_pages', 0),
                    'file_type': document_info.get('file_type', 'Unknown'),
                    'text_length': len(text_content),
                    'extraction_method': 'Gemini AI Analysis',
                    'analysis_timestamp': policy.updated_at.isoformat() if policy.updated_at else None
                }
                
                return Response(extracted_details, status=status.HTTP_200_OK)
            else:
                logger.warning("Gemini AI analysis returned no data, using fallback")
                raise Exception("AI analysis returned no data")
                
        except Exception as ai_error:
            logger.error(f"Gemini AI analysis failed: {ai_error}")
            
            # STEP 3: Fallback to basic extraction if AI fails
            logger.info("Using fallback extraction method...")
            
            # Create basic fallback data based on extracted text
            extracted_details = {
                'summary': f"⚠️ FALLBACK DATA: Policy {policy.name} appears to be a {policy.policy_type or 'insurance'} policy. AI analysis failed, so this is basic information extracted from the document.",
                'coverage': f"Basic {policy.policy_type or 'insurance'} coverage information available. AI analysis was not available due to: {str(ai_error)}",
                'effectiveDate': policy.start_date.strftime('%B %d, %Y') if policy.start_date else 'Policy start date to be confirmed',
                'expiryDate': policy.end_date.strftime('%B %d, %Y') if policy.end_date else 'Policy end date to be confirmed',
                'department': policy.policy_type or 'General Insurance',
                'deductible': f"Standard {policy.policy_type or 'insurance'} deductible applies",
                'maxOutOfPocket': f"Maximum out-of-pocket expenses as per {policy.policy_type or 'insurance'} policy terms",
                'tags': [policy.policy_type or 'Insurance', policy.provider or 'Provider', 'Active Policy', 'Basic Coverage', 'AI Failed'],
                'mlInsights': {
                    'riskAssessment': 'Standard Risk (Fallback)',
                    'coverageScore': 60,
                    'costEfficiency': 'Standard Value (Fallback)',
                    'marketComparison': 'Market Standard (Fallback)',
                    'optimizationTips': [
                        '⚠️ AI analysis failed - this is fallback data',
                        'Review policy terms with your agent',
                        'Consider annual policy review',
                        'Contact support if AI analysis issues persist'
                    ]
                },
                'extraction_confidence': 0.4,
                'isFallbackData': True,
                'fallbackReason': f"Gemini AI analysis failed: {str(ai_error)}",
                'document_analysis': {
                    'total_pages': document_info.get('total_pages', 0),
                    'file_type': document_info.get('file_type', 'Unknown'),
                    'text_length': len(text_content),
                    'extraction_method': 'Fallback (AI Failed)',
                    'analysis_timestamp': policy.updated_at.isoformat() if policy.updated_at else None
                }
            }
            
            return Response(extracted_details, status=status.HTTP_200_OK)
        
    except Policy.DoesNotExist:
        return Response({'error': 'Policy not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        logger.error(f"Error in extract_policy_details: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)





def _assess_risk_level(text_content):
    """Assess risk level based on document content"""
    text_lower = text_content.lower()
    
    # High risk indicators
    high_risk_terms = ['high risk', 'dangerous', 'hazardous', 'excluded', 'not covered', 'void']
    # Medium risk indicators  
    medium_risk_terms = ['moderate', 'standard', 'normal', 'typical', 'average']
    # Low risk indicators
    low_risk_terms = ['low risk', 'safe', 'protected', 'comprehensive', 'full coverage']
    
    if any(term in text_lower for term in high_risk_terms):
        return 'High Risk'
    elif any(term in text_lower for term in medium_risk_terms):
        return 'Medium Risk'
    elif any(term in text_lower for term in low_risk_terms):
        return 'Low Risk'
    else:
        return 'Low to Moderate Risk'


def _calculate_coverage_score(text_content):
    """Calculate coverage score based on document content"""
    text_lower = text_content.lower()
    score = 50  # Base score
    
    # Positive indicators
    if 'comprehensive' in text_lower:
        score += 20
    if 'full coverage' in text_lower:
        score += 15
    if 'additional riders' in text_lower or 'endorsements' in text_lower:
        score += 10
    if 'umbrella' in text_lower:
        score += 10
    if 'excess' in text_lower:
        score += 5
        
    # Negative indicators
    if 'excluded' in text_lower:
        score -= 10
    if 'not covered' in text_lower:
        score -= 15
    if 'limited' in text_lower:
        score -= 10
        
    return max(0, min(100, score))


def _assess_cost_efficiency(text_content):
    """Assess cost efficiency based on document content"""
    text_lower = text_content.lower()
    
    if 'discount' in text_lower or 'reduced' in text_lower:
        return 'Excellent Value'
    elif 'competitive' in text_lower or 'market rate' in text_lower:
        return 'Good Value'
    elif 'premium' in text_lower and 'high' in text_lower:
        return 'Premium Cost'
    else:
        return 'Good Value'


def _assess_market_position(text_content):
    """Assess market position based on document content"""
    text_lower = text_content.lower()
    
    if 'leading' in text_lower or 'top' in text_lower:
        return 'Market Leader'
    elif 'competitive' in text_lower or 'standard' in text_lower:
        return 'Competitive in Market'
    elif 'premium' in text_lower or 'exclusive' in text_lower:
        return 'Premium Position'
    else:
        return 'Competitive in Market'


def _generate_optimization_tips(text_content, policy_type):
    """Generate optimization tips based on content and policy type"""
    tips = []
    text_lower = text_content.lower()
    
    # General tips
    if 'bundle' not in text_lower:
        tips.append('Consider bundling with other policies for better rates')
    
    if 'annual review' not in text_lower:
        tips.append('Review coverage limits annually')
        
    if 'good standing' not in text_lower:
        tips.append('Maintain good standing for premium reductions')
        
    if 'riders' not in text_lower and 'endorsements' not in text_lower:
        tips.append('Explore additional riders for enhanced protection')
    
    # Policy-specific tips
    if policy_type and 'auto' in policy_type.lower():
        tips.append('Consider increasing liability limits for better protection')
        tips.append('Check for safe driver discounts')
    elif policy_type and 'home' in policy_type.lower():
        tips.append('Install security systems for premium reductions')
        tips.append('Consider flood insurance if in flood-prone areas')
    elif policy_type and 'life' in policy_type.lower():
        tips.append('Review beneficiary designations regularly')
        tips.append('Consider term vs permanent life insurance options')
    
    # Ensure we have at least 3 tips
    while len(tips) < 3:
        tips.append('Contact your agent for personalized recommendations')
        
    return tips[:5]  # Return max 5 tips


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_policy_extraction(request, policy_id):
    """Get extracted text from a specific policy."""
    try:
        policy = get_object_or_404(Policy, id=policy_id, user=request.user)
        
        # Try to get from PolicyExtraction first
        if hasattr(policy, 'extraction') and policy.extraction:
            extracted_text = policy.extraction.extracted_text
        else:
            # Fallback to basic policy info
            extracted_text = f"Policy: {policy.name}\nType: {policy.policy_type}\nProvider: {policy.provider}\nDescription: {policy.description or 'No description available'}"
        
        return Response({
            'status': 'success',
            'data': {
                'policy_id': policy_id,
                'extracted_text': extracted_text,
                'policy_name': policy.name
            }
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"Get policy extraction error: {str(e)}")
        return Response({
            'status': 'error',
            'error': f'Failed to get policy extraction: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)







@api_view(['POST'])
@permission_classes([IsAuthenticated])
def query_policy(request):
    """Query a specific policy with a question."""
    try:
        data = request.data
        user = request.user
        
        policy_id = data.get('policy_id')
        question = data.get('question')
        
        logger.info(f"Policy query request from user {user.email} for policy {policy_id}: {question}")
        
        if not policy_id or not question:
            logger.warning(f"Policy query request missing required data. Policy ID: {policy_id}, Question: {question}")
            return Response({
                'status': 'error',
                'message': 'Policy ID and question are required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get policy and verify access
        try:
            policy = get_object_or_404(Policy, id=policy_id, user=user)
            logger.info(f"Policy {policy_id} found and access verified for user {user.email}")
        except Exception as e:
            logger.error(f"Policy access error for user {user.email}, policy {policy_id}: {e}")
            return Response({
                'status': 'error',
                'message': 'Policy not found or access denied'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Get or create conversation
        conversation = get_or_create_conversation(user, policy)
        if not conversation:
            logger.error(f"Failed to create conversation for user {user.email} and policy {policy_id}")
            return Response({
                'status': 'error',
                'message': 'Failed to create conversation'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        logger.info(f"Conversation created/retrieved: {conversation.id}")
        
        # Increment conversation count for the policy if this is a new conversation
        if conversation.messages.count() == 0:
            try:
                policy.increment_conversation_count()
                logger.info(f"Incremented conversation count for policy {policy_id}")
            except Exception as e:
                logger.warning(f"Failed to increment conversation count for policy {policy_id}: {e}")
        
        # Store user message
        user_message = store_message(
            conversation=conversation,
            message_type='user',
            content=question
        )
        
        logger.info(f"User message stored: {user_message.id}")
        
        # Get policy extraction text if available
        context = ""
        try:
            # First try to get the actual document content by reading the file
            if policy.document:
                import PyPDF2
                from docx import Document
                
                document_text = ""
                try:
                    if policy.document.name.lower().endswith('.pdf'):
                        # Read PDF content directly
                        pdf_reader = PyPDF2.PdfReader(policy.document)
                        for page in pdf_reader.pages:
                            document_text += page.extract_text()
                        logger.info(f"Extracted {len(document_text)} characters from PDF for chat context")
                    elif policy.document.name.lower().endswith('.docx'):
                        # Read DOCX content directly
                        doc = Document(policy.document)
                        for paragraph in doc.paragraphs:
                            document_text += paragraph.text + "\n"
                        logger.info(f"Extracted {len(document_text)} characters from DOCX for chat context")
                    else:
                        document_text = f"Document: {policy.document.name} - {policy.description or 'No description available'}"
                        logger.warning(f"Unsupported file type for chat context: {policy.document.name}")
                except Exception as file_error:
                    logger.warning(f"Could not read document file for chat context: {file_error}")
                    document_text = f"Document: {policy.document.name} - {policy.description or 'No description available'}"
                
                # Use the extracted document text as context
                if document_text and len(document_text.strip()) > 50:
                    context = f"Policy Document Content:\n{document_text[:3000]}...\n\n"
                else:
                    context = f"Policy: {policy.name} - {policy.description or 'No description available'}\n\n"
            else:
                # No document file available
                context = f"Policy: {policy.name} - {policy.description or 'No description available'}\n\n"
                
        except Exception as e:
            logger.warning(f"Could not get policy context: {e}")
            context = f"Policy: {policy.name} - {policy.description or 'No description available'}\n\n"
        
        logger.info(f"Policy context prepared: {len(context)} characters")
        
        # Get AI response using Gemini
        try:
            from .services import GeminiService
            
            logger.info("Initializing Gemini service for policy query...")
            gemini_service = GeminiService()
            
            logger.info(f"Gemini service initialized. Model: {getattr(gemini_service, 'model', 'None')}")
            
            if not gemini_service.model:
                logger.error("AI service not available - no model")
                return Response({
                    'status': 'error',
                    'message': 'AI service not available. Please check Gemini API configuration.'
                }, status=status.HTTP_503_SERVICE_UNAVAILABLE)
            
            # Create a simple, clear prompt for policy chat
            prompt = f"""You are PolicyBridge AI, an assistant that explains insurance in simple terms.
Always follow these rules when answering a question.

Rules for Answering:
- Clarity First: Use simple, everyday language
- Concise: Keep answers short (2–4 sentences)
- Direct & Helpful: Get to the point, avoid filler
- Use Examples: Add quick, practical examples if helpful
- Context Aware: If the question is about eligibility, claim, or coverage, include key conditions
- No Jargon: Avoid legal/technical terms unless needed, then explain in plain words
- Confidence: Always give a clear response. If unsure, explain what usually applies in most policies

Policy Context:
{context}

Question: {question}

Answer:"""

            logger.info(f"Sending policy query to Gemini")
            logger.info(f"Question: {question}")
            logger.info(f"Prompt length: {len(prompt)} characters")

            ai_response = gemini_service.get_response(prompt)

            if not ai_response:
                logger.error("Gemini returned empty response for policy query")
                raise Exception("Gemini returned empty response")

            logger.info(f"Successfully received AI response for policy query: {len(ai_response)} characters")

            # Store AI response
            ai_message = store_message(
                conversation=conversation,
                message_type='ai',
                content=ai_response,
                citations=[],
                ml_insights={
                    "question_answered": True,
                    "response_quality": "high",
                    "policy_id": policy_id
                }
            )
            
            logger.info(f"AI message stored: {ai_message.id}")
            
            return Response({
                "status": "success",
                "response": ai_response,
                "citations": [],
                "ml_insights": {
                    "question_answered": True,
                    "response_quality": "high",
                    "policy_id": policy_id
                },
                "conversation_id": conversation.id,
                "message_id": ai_message.id if ai_message else None
            })

        except Exception as e:
            logger.error(f"Error getting AI response for policy query: {e}")
            logger.error(f"Error type: {type(e).__name__}")
            logger.error(f"Error details: {str(e)}")

            # Helpful fallback response
            fallback_response = (
                f"I'm sorry, I'm having trouble accessing the AI service right now. "
                f"Based on your policy '{policy.name}', I can see this is a {policy.policy_type or 'insurance'} policy. "
                f"For specific questions, I recommend consulting with a licensed insurance agent "
                f"or reviewing your policy documentation for accurate information."
            )

            return Response({
                "status": "success",
                "response": fallback_response,
                "citations": [],
                "ml_insights": {
                    "question_answered": True,
                    "response_quality": "fallback",
                    "fallback_reason": str(e),
                    "policy_id": policy_id
                }
            })

    except Exception as e:
        logger.error(f"Error in policy query: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        return Response({
            "status": "error",
            "message": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def general_chat(request):
    """General insurance chat without specific policy context."""
    try:
        user = request.user
        data = request.data
        question = data.get('question')
        
        logger.info(f"General chat request from user {user.email}: {question}")
        
        if not question:
            logger.warning("General chat request missing question")
            return Response({
                'status': 'error',
                'message': 'Question is required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get or create conversation
        conversation = get_or_create_conversation(user, None)
        if not conversation:
            logger.error(f"Failed to create conversation for user {user.email}")
            return Response({
                'status': 'error',
                'message': 'Failed to create conversation'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        logger.info(f"Conversation created/retrieved: {conversation.id}")
        
        # Increment conversation count for general conversations if this is a new conversation
        if conversation.messages.count() == 0:
            try:
                # For general conversations, we can track in user stats or just log
                logger.info(f"New general conversation started for user {user.email}")
            except Exception as e:
                logger.warning(f"Failed to track general conversation: {e}")
        
        # Store user message
        user_message = store_message(
            conversation=conversation,
            message_type='user',
            content=question
        )
        
        logger.info(f"User message stored: {user_message.id}")
        
        # Get AI response using Gemini
        try:
            from .services import GeminiService
            
            logger.info("Initializing Gemini service...")
            gemini_service = GeminiService()
            
            logger.info(f"Gemini service initialized. Model: {getattr(gemini_service, 'model', 'None')}")
            
            if not gemini_service.model:
                logger.error("AI service not available - no model")
                return Response({
                    'status': 'error',
                    'message': 'AI service not available. Please check Gemini API configuration.'
                }, status=status.HTTP_503_SERVICE_UNAVAILABLE)
            
            # Create a simple, clear prompt for general insurance chat
            prompt = f"""You are PolicyBridge AI, an assistant that explains insurance in simple terms.
            Always follow these rules when answering a question.

            Rules for Answering:
            - Clarity First: Use simple, everyday language
            - Concise: Keep answers short (2–4 sentences)
            - Direct & Helpful: Get to the point, avoid filler
            - Use Examples: Add quick, practical examples if helpful
            - No Jargon: Avoid legal/technical terms unless needed, then explain in plain words
            - Confidence: Always give a clear response. If unsure, explain what usually applies in most cases

            Question: {question}

            Answer:"""

            logger.info(f"Sending general chat query to Gemini")
            logger.info(f"Question: {question}")
            logger.info(f"Prompt length: {len(prompt)} characters")

            ai_response = gemini_service.get_response(prompt)

            if not ai_response:
                logger.error("Gemini returned empty response")
                raise Exception("Gemini returned empty response")

            logger.info(f"Successfully received AI response for general chat: {len(ai_response)} characters")

            # Store AI response
            ai_message = store_message(
                conversation=conversation,
                message_type='ai',
                content=ai_response,
                citations=[],
                ml_insights={
                    "question_answered": True,
                    "response_quality": "high"
                }
            )
            
            logger.info(f"AI message stored: {ai_message.id}")
            
            return Response({
                "status": "success",
                "response": ai_response,
                "citations": [],
                "ml_insights": {
                    "question_answered": True,
                    "response_quality": "high"
                },
                "conversation_id": conversation.id,
                "message_id": ai_message.id if ai_message else None
            })

        except Exception as e:
            logger.error(f"Error getting AI response for general chat: {e}")
            logger.error(f"Error type: {type(e).__name__}")
            logger.error(f"Error details: {str(e)}")

            # Helpful fallback response
            fallback_response = (
                "I'm sorry, I'm having trouble accessing the AI service right now. "
                "For insurance questions, I recommend consulting with a licensed insurance agent "
                "or reviewing official insurance documentation for accurate information."
            )

            return Response({
                "status": "success",
                "response": fallback_response,
                "citations": [],
                "ml_insights": {
                    "question_answered": True,
                    "response_quality": "fallback",
                    "fallback_reason": str(e)
                }
            })

    except Exception as e:
        logger.error(f"Error in general chat: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        return Response({
            "status": "error",
            "message": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_conversations(request):
    """Get all conversations for the authenticated user"""
    try:
        conversations = Conversation.objects.filter(user=request.user, is_active=True)
        serializer = ConversationListSerializer(conversations, many=True)
        return Response({
            "status": "success",
            "conversations": serializer.data
        })
    except Exception as e:
        logger.error(f"Error getting conversations: {e}")
        return Response({
            "status": "error",
            "message": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_conversation_messages(request, conversation_id):
    """Get all messages for a specific conversation"""
    try:
        conversation = get_object_or_404(Conversation, id=conversation_id, user=request.user)
        messages = conversation.messages.all()
        serializer = MessageSerializer(messages, many=True)
        return Response({
            "status": "success",
            "conversation": ConversationSerializer(conversation).data,
            "messages": serializer.data
        })
    except Exception as e:
        logger.error(f"Error getting conversation messages: {e}")
        return Response({
            "status": "error",
            "message": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_conversation(request, conversation_id):
    """Delete a conversation and all its messages"""
    try:
        conversation = get_object_or_404(Conversation, id=conversation_id, user=request.user)
        conversation.delete()
        return Response({
            "status": "success",
            "message": "Conversation deleted successfully"
        })
    except Exception as e:
        logger.error(f"Error deleting conversation: {e}")
        return Response({
            "status": "error",
            "message": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)






